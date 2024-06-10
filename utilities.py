from sympy import symbols, Eq, solve
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import random
from bs4 import BeautifulSoup

# subscript dictionary
subscript_dict = {
    "1": "₁",
    "2": "₂",
    "3": "₃",
    "4": "₄",
    "5": "₅",
    "6": "₆",
    "7": "₇",
    "8": "₈",
    "9": "₉",
}


def process_and_save_subscripts(filename):
    with open(filename, "r", encoding="utf-8") as file:
        file_text = file.read()
    print(file_text)
    file_text = subs_tag_html(file_text)
    print(file_text)
    # Save the modified HTML content to a new file
    with open(filename, "w", encoding="utf-8") as f:
        f.write(file_text)


def get_reactions_list(filename="BalancedReactions.txt"):
    with open(filename, "r", encoding="utf-8") as file:
        balancedEquations = file.readlines()
    return balancedEquations


# Function to subscript numbers after elements
def subscript_equation(equation):
    # Replace each number after an element with its subscript equivalent
    for number, subscript in subscript_dict.items():
        equation = re.sub(r"(?<=[A-Za-z])" + number, subscript, equation)
    return equation


# Function to subscript numbers after elements
def subs_tag_html(html_string):
    # Replace each number after an element with its subscript equivalent
    for number, subscript in subscript_dict.items():
        html_string = re.sub(subscript, f"<sub>{number}</sub>", html_string)

    html_string = re.sub("→", "&rarr;", html_string)
    return html_string


def multiply_prefix_numbers(match, multiplyBy):
    number = match.group(1)
    return str(int(number) * multiplyBy)


def space_out_equations(equations):
    # space out equations
    for i in range(len(equations)):
        equations[i] = equations[i].replace("+", " + ").replace("→", " → ")
        equations[i] = equations[i].replace("\u200b", "")
    return equations


def create_equation_document(equations, teacher_answer_sheet):

    doc = Document()
    if teacher_answer_sheet:
        filename = "GeneratedDocuments/AnswerDocument.docx"
        doc.add_paragraph(f"TEACHER ANSWER SHEET")
        doc.add_paragraph(f"Balance the following chemical equations:")
    else:
        doc.add_paragraph(f"Balance the following chemical equations:")
        filename = "GeneratedDocuments/worksheetGenerated.docx"

    for equation in equations:
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = para.add_run(equation)
        run.add_break()
        run.font.size = Pt(14)

    doc.save(filename)
    return doc


def generate_random_equations(equations):

    # get equations
    for i in range(len(equations)):

        eq = equations[i]

        # subscript and remove invisible values
        eq = subscript_equation(eq)
        eq = re.sub(r"[^\w+→ ]", "", eq)

        # multiple all prefix values by a random integer 1-10
        multiply_by = random.randrange(1, 5)
        lambda_multiply_function = lambda match: multiply_prefix_numbers(
            match, multiply_by
        )

        eq = re.sub(r"(\d+)(?=[A-Za-z])", lambda_multiply_function, eq)

        if multiply_by > 1:
            # Prepend 'multiply_by' to elements that are not preceded by a number (equivalent of 1)
            eq = re.sub(
                r"((?<=^)|(?<=\s)|(?<=\+)|(?<=\→))([A-Za-z])",
                str(multiply_by) + r"\2",
                eq,
            )

        # if hide_random_values > 0:

        #     # Find all matches
        #     matches = list(re.finditer(r"(\d+)(?=[A-Za-z])", eq))

        #     if matches:
        #         for _ in range(hide_random_values):
        #             # Choose a random match
        #             match = random.choice(matches)

        #             # Replace the random match with '_'
        #             eq = eq[: match.start()] + "__" + eq[match.end() :]
        #             matches.remove(match)
        # else:
        #     # replace all prefix numbers with '__'
        #     if hide_all_values:
        #         eq = re.sub(
        #             r"((?<=^)|(?<=\s)|(?<=\+)|(?<=\→))([A-Za-z])", str(1) + r"\2", eq
        #         )
        #         eq = re.sub(r"(\d+)(?=[A-Za-z])", "__", eq)

        #     # replace prefix numbers with '__'
        #     if hide_one_value:
        #         eq = re.sub(r"(\d+)(?=[A-Za-z])", "__", eq, count=1)
        equations[i] = eq

    return equations


def hide_numbers_in_equations(
    equations, hide_all_values=False, hide_one_value=False, hide_random_values=0
):

    # get equations
    for i in range(len(equations)):

        eq = equations[i]

        if hide_random_values > 0:

            # Find all matches
            matches = list(re.finditer(r"(\d+)(?=[A-Za-z])", eq))

            for _ in range(hide_random_values):
                if matches:
                    # Choose a random match
                    match = random.choice(matches)

                    # Replace the random match with '_'
                    eq = eq[: match.start()] + "__" + eq[match.end() :]
                    matches.remove(match)
        else:
            # replace all prefix numbers with '__'
            if hide_all_values:
                eq = re.sub(
                    r"((?<=^)|(?<=\s)|(?<=\+)|(?<=\→))([A-Za-z])", str(1) + r"\2", eq
                )
                eq = re.sub(r"(\d+)(?=[A-Za-z])", "__", eq)

            # replace prefix numbers with '__'
            if hide_one_value:

                eq = re.sub(r"(\d+)(?=[A-Za-z])", "__", eq, count=1)

        equations[i] = eq

    return equations


def create_basic_document(
    numberEquations, hide_all_values, hide_one_value, hide_random_values
):
    balancedEquations = get_reactions_list()
    spaced_out_equations = space_out_equations(balancedEquations)
    randomSelection = random.choices(spaced_out_equations, k=numberEquations)
    equations = generate_random_equations(randomSelection)

    # create teacher document
    create_equation_document(equations, True)

    # create student document with hidden number
    equations = hide_numbers_in_equations(
        equations,
        hide_all_values=hide_all_values,
        hide_one_value=hide_one_value,
        hide_random_values=hide_random_values,
    )
    doc = create_equation_document(equations, False)
    return doc
